# Validate Tables
# @author: Junsoo Park
#
# This program will take two excel files and compare their common columns to see if they match.
# It assumes that the first table is a subset of the second (i.e. the first table is a view of the second).
# It asks for user input and allows the user to specify the actual column names if they are not matching already.
# It produces a summary report of the validation.

import pandas as pd
from docx import Document
from docx.shared import Inches

# Reading Excel files and converting them to data frames
file1 = pd.read_excel('path/to/file1.xlsx')
file2 = pd.read_excel('path/to/file2.xlsx')

# Prompting user for input on column names
all_cols_same = input("Are all column headers the same in both files? (y/n) ")
if all_cols_same.lower() == 'y':
    common_cols = file1.columns
else:
    common_cols_input = input("Please enter the common column names, separated by commas: ")
    common_cols = common_cols_input.split(',')
    common_cols = [col.strip() for col in common_cols]

    # Subsetting data frames to include only the common columns
    file1 = file1[common_cols]
    file2 = file2[common_cols]

# Creating a new Word document
document = Document()

# Adding a title to the document
document.add_heading('Comparison Report', level=0)

# Counters for rows and columns that did not match
row_mismatches = 0
col_mismatches = 0

# Looping through rows and comparing values
for index, row in file1.iterrows():
    match = True
    for col in common_cols:
        if pd.isna(row[col]) or pd.isna(file2.iloc[index][col]):
            if pd.isna(row[col]) and pd.isna(file2.iloc[index][col]):
                continue
            else:
                match = False
                col_mismatches += 1
                break
        elif row[col] != file2.iloc[index][col]:
            match = False
            col_mismatches += 1
            break
    if not match:
        row_mismatches += 1
        # Adding a table to the document with the row and column information
        table = document.add_table(rows=2, cols=len(common_cols))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(common_cols):
            hdr_cells[i].text = col
        row_cells = table.rows[1].cells
        for i, col in enumerate(common_cols):
            row_cells[i].text = str(row[col])
        for i, col in enumerate(common_cols):
            row_cells[i].text += '\n'
            row_cells[i].text += str(file2.iloc[index][col])

# Adding summary information to the document
document.add_heading('Summary', level=1)
document.add_paragraph(f"Total number of rows checked: {len(file1)}")
document.add_paragraph(f"Total number of columns checked: {len(common_cols)}")
document.add_paragraph(f"Number of rows that did not match: {row_mismatches}")
document.add_paragraph(f"Number of column values that did not match: {col_mismatches}")

# Saving the document
document.save('Comparison Report.docx')